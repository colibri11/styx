"""DOCX parser через python-docx (волна 28).

Pure-Python чтение Word XML внутри .docx zip-контейнера. Сохраняет
параграфы + содержимое таблиц. Headers/footers — не вытаскиваем
(out of scope волны 28).
"""

from __future__ import annotations

from pathlib import Path

from styx.engine.document_parsers._types import ParsedDocument


def parse_docx(path: Path) -> ParsedDocument:
    """Извлечь текст из .docx.

    Raises:
        ValueError: malformed DOCX (python-docx не распарсил).
    """
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = Document(str(path))
    except PackageNotFoundError as exc:
        raise ValueError(f"DOCX parse error: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                (cell.text or "").strip() for cell in row.cells
            ).strip(" |")
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()

    return ParsedDocument(
        text=text,
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
    )
