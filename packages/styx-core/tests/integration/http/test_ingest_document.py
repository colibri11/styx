"""POST /ingest_document integration: TestClient + real Postgres +
Ollama (волна 28).

Plugin-side path-mode: core читает абсолютный path с диска (D1), парсит
(pypdf/python-docx/openpyxl/builtin), режет на chunks, embed'ит,
INSERT'ит document + chunks. tail-memory НЕ создаётся (D5).

Tested:
- Все 5 форматов (PDF/DOCX/XLSX/MD/TXT) → 200, document + chunks
  записаны, embedding'и не NULL, tail-memory нет.
- Идемпотентность (SHA256 file_bytes) — повторный POST → deduplicated.
- 422 cases: relative path, file not found, unsupported extension,
  mime mismatch, empty text (blank PDF), file size exceeds limit.
- Whitelist enforcement (STYX_INGEST_DOC_ROOTS).
- Hash override через request.
"""

from __future__ import annotations

import os
import uuid
from dataclasses import replace
from pathlib import Path

import psycopg
import pytest
from fastapi.testclient import TestClient

from styx.config import StyxConfig, load as load_config
from styx.http import registry
from styx.http.app import create_app
from styx.providers.memory import StyxMemoryCore
from styx.storage import migrate


pytestmark = pytest.mark.skipif(
    not os.environ.get("STYX_TEST_DATABASE_URL"),
    reason="STYX_TEST_DATABASE_URL не задан — integration tests skip",
)


FIXTURE_DIR = Path(__file__).parent.parent.parent / "fixtures" / "documents"


# ── Stack fixture ───────────────────────────────────────────────────


@pytest.fixture
def stack(clean_db: str):
    migrate.run(clean_db)

    cfg: StyxConfig = load_config()
    cfg = replace(cfg, database_url=clean_db, http_token=None)

    agent = "alpha"
    core = StyxMemoryCore(agent_id=agent)
    core._config = cfg
    core.initialize(session_id=str(uuid.uuid4()), agent_identity=agent)
    registry.reset_all()
    registry.register(agent_id=agent, core=core)

    app = create_app(cfg)
    client = TestClient(app)

    yield client, agent, clean_db, core, cfg

    core.shutdown()
    registry.reset_all()


# ── Binary fixture builders ─────────────────────────────────────────


@pytest.fixture
def pdf_with_text(tmp_path: Path) -> Path:
    """PDF с реальным текстом через pypdf low-level API."""
    from pypdf import PdfWriter
    from pypdf.generic import (
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
    )

    writer = PdfWriter()
    page = writer.add_blank_page(width=595, height=842)
    stream = DecodedStreamObject()
    stream.set_data(
        b"BT /F1 24 Tf 100 700 Td (Wave 28 ingest pipeline) Tj ET"
    )
    page[NameObject("/Contents")] = stream
    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    resources = page.get("/Resources", DictionaryObject())
    fonts = resources.get("/Font", DictionaryObject())
    fonts[NameObject("/F1")] = font
    resources[NameObject("/Font")] = fonts
    page[NameObject("/Resources")] = resources

    target = tmp_path / "sample.pdf"
    with target.open("wb") as fh:
        writer.write(fh)
    return target


@pytest.fixture
def pdf_blank(tmp_path: Path) -> Path:
    """Image-only PDF — blank page, extract_text возвращает empty."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    target = tmp_path / "blank.pdf"
    with target.open("wb") as fh:
        writer.write(fh)
    return target


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Параграф 1 — содержит русские слова.")
    doc.add_paragraph("Paragraph 2 — and English words.")
    target = tmp_path / "sample.docx"
    doc.save(str(target))
    return target


@pytest.fixture
def xlsx_file(tmp_path: Path) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Data"
    ws.append(["Name", "Value"])
    ws.append(["one", 42])
    target = tmp_path / "sample.xlsx"
    wb.save(str(target))
    return target


# ── Postgres helpers ────────────────────────────────────────────────


def _document_columns(dsn: str, document_id: str) -> dict:
    with psycopg.connect(dsn) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT file_path, original_name, mime_type, "
                "       size_bytes, char_count, visibility, source, "
                "       content_hash, metadata, summary, agent_id "
                "FROM documents WHERE id = %s",
                (document_id,),
            )
            row = cur.fetchone()
    cols = [
        "file_path", "original_name", "mime_type", "size_bytes",
        "char_count", "visibility", "source", "content_hash",
        "metadata", "summary", "agent_id",
    ]
    return dict(zip(cols, row))


def _chunks_count(dsn: str, document_id: str) -> tuple[int, int]:
    """Возвращает (total_chunks, with_embedding)."""
    with psycopg.connect(dsn) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT count(*), count(embedding) FROM chunks "
                "WHERE document_id = %s",
                (document_id,),
            )
            row = cur.fetchone()
    return (row[0], row[1])


def _tail_memory_count(dsn: str, agent: str) -> int:
    """Любые ряды в memories у agent'а (для проверки что file-ingest НЕ
    создал tail-memory)."""
    with psycopg.connect(dsn) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT count(*) FROM memories WHERE agent_id = %s",
                (agent,),
            )
            row = cur.fetchone()
    return row[0]


# ── Happy paths ─────────────────────────────────────────────────────


def test_ingest_plaintext(stack) -> None:
    client, agent, dsn, _, _ = stack
    target = FIXTURE_DIR / "sample.txt"
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(target)},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["deduplicated"] is False
    assert body["mime_type"] == "text/plain"
    assert body["original_name"] == "sample.txt"
    assert body["chunks_count"] >= 1
    assert len(body["content_hash"]) == 64

    cols = _document_columns(dsn, body["document_id"])
    assert cols["agent_id"] == agent
    assert cols["source"] == "ingest_document"
    assert cols["mime_type"] == "text/plain"
    assert cols["file_path"] == str(target.resolve())
    assert cols["original_name"] == "sample.txt"
    assert cols["size_bytes"] > 0
    assert cols["content_hash"] == body["content_hash"]

    total, with_emb = _chunks_count(dsn, body["document_id"])
    assert total == body["chunks_count"]
    assert with_emb == total  # все chunk'и с embedding'ами

    # tail-memory НЕ создаётся (D5).
    assert _tail_memory_count(dsn, agent) == 0


def test_ingest_markdown(stack) -> None:
    client, agent, dsn, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(FIXTURE_DIR / "sample.md")},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["mime_type"] == "text/markdown"
    cols = _document_columns(dsn, body["document_id"])
    assert "line_count" in cols["metadata"]


def test_ingest_pdf(stack, pdf_with_text: Path) -> None:
    client, agent, dsn, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(pdf_with_text)},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["mime_type"] == "application/pdf"
    assert "Wave 28 ingest pipeline" not in body  # text не в response
    cols = _document_columns(dsn, body["document_id"])
    assert "page_count" in cols["metadata"]


def test_ingest_docx(stack, docx_file: Path) -> None:
    client, agent, dsn, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(docx_file)},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert (
        body["mime_type"]
        == "application/vnd.openxmlformats-officedocument."
           "wordprocessingml.document"
    )
    cols = _document_columns(dsn, body["document_id"])
    assert "paragraph_count" in cols["metadata"]


def test_ingest_xlsx(stack, xlsx_file: Path) -> None:
    client, agent, dsn, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(xlsx_file)},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert (
        body["mime_type"]
        == "application/vnd.openxmlformats-officedocument."
           "spreadsheetml.sheet"
    )
    cols = _document_columns(dsn, body["document_id"])
    assert "sheet_names" in cols["metadata"]
    assert cols["metadata"]["sheet_names"] == ["Data"]


# ── Idempotency ─────────────────────────────────────────────────────


def test_ingest_idempotent_same_file(stack) -> None:
    """Повторный ingest того же файла → deduplicated=True, тот же
    document_id, новых INSERT'ов нет."""
    client, agent, dsn, _, _ = stack
    target = FIXTURE_DIR / "sample.txt"
    first = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(target)},
    )
    second = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(target)},
    )
    assert first.status_code == 200 and second.status_code == 200
    fb, sb = first.json(), second.json()
    assert fb["document_id"] == sb["document_id"]
    assert fb["deduplicated"] is False
    assert sb["deduplicated"] is True
    assert sb["chunks_count"] == 0  # без INSERT'ов


def test_ingest_explicit_hash_override(stack) -> None:
    client, agent, dsn, _, _ = stack
    explicit = "deadbeef" * 8
    resp = client.post(
        "/ingest_document",
        json={
            "agent_id": agent,
            "path": str(FIXTURE_DIR / "sample.txt"),
            "content_hash": explicit,
        },
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["content_hash"] == explicit
    cols = _document_columns(dsn, body["document_id"])
    assert cols["content_hash"] == explicit


def test_ingest_source_ref_and_visibility(stack) -> None:
    client, agent, dsn, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={
            "agent_id": agent,
            "path": str(FIXTURE_DIR / "sample.md"),
            "source_ref": "upload-42",
            "visibility": "private",
            "metadata": {"upload_source": "openclaw"},
        },
    )
    assert resp.status_code == 200, resp.text
    cols = _document_columns(dsn, resp.json()["document_id"])
    assert cols["source"] == "ingest_document"
    # Note: source_ref / visibility — отдельные колонки (миграция 0007).
    with psycopg.connect(dsn) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT source_ref, visibility FROM documents WHERE id = %s",
                (resp.json()["document_id"],),
            )
            row = cur.fetchone()
    assert row == ("upload-42", "private")
    # User metadata в documents.metadata JSONB.
    assert cols["metadata"]["upload_source"] == "openclaw"


# ── 422 cases ───────────────────────────────────────────────────────


def test_ingest_422_relative_path(stack) -> None:
    client, agent, _, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": "relative/file.txt"},
    )
    assert resp.status_code == 422
    assert "must be absolute" in resp.json()["detail"]


def test_ingest_422_file_not_found(stack, tmp_path: Path) -> None:
    client, agent, _, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(tmp_path / "missing.pdf")},
    )
    assert resp.status_code == 422
    assert "file not found" in resp.json()["detail"]


def test_ingest_422_unsupported_extension(
    stack, tmp_path: Path
) -> None:
    client, agent, _, _, _ = stack
    decoy = tmp_path / "x.pptx"
    decoy.write_text("not really pptx")
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(decoy)},
    )
    assert resp.status_code == 422
    assert "unsupported extension" in resp.json()["detail"]


def test_ingest_422_mime_mismatch(stack, tmp_path: Path) -> None:
    client, agent, _, _, _ = stack
    decoy = tmp_path / "fake.pdf"
    decoy.write_bytes(b"PK\x03\x04not a real pdf")
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(decoy)},
    )
    assert resp.status_code == 422
    assert "mime mismatch" in resp.json()["detail"]


def test_ingest_422_empty_document(stack, pdf_blank: Path) -> None:
    """Blank PDF (image-only) — empty text → 422."""
    client, agent, _, _, _ = stack
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(pdf_blank)},
    )
    assert resp.status_code == 422
    assert "empty document" in resp.json()["detail"]


def test_ingest_422_empty_text_file(stack, tmp_path: Path) -> None:
    client, agent, _, _, _ = stack
    target = tmp_path / "empty.txt"
    target.write_text("")
    resp = client.post(
        "/ingest_document",
        json={"agent_id": agent, "path": str(target)},
    )
    assert resp.status_code == 422
    assert "empty document" in resp.json()["detail"]


# ── Disabled toggle ─────────────────────────────────────────────────


def test_ingest_disabled_returns_503(clean_db: str, tmp_path: Path) -> None:
    """ingest_doc_enabled=False → 503."""
    migrate.run(clean_db)
    cfg = load_config()
    cfg = replace(
        cfg,
        database_url=clean_db,
        http_token=None,
        ingest_doc_enabled=False,
    )

    agent = "alpha"
    core = StyxMemoryCore(agent_id=agent)
    core._config = cfg
    core.initialize(session_id=str(uuid.uuid4()), agent_identity=agent)
    # initialize перезатёр _config — ставим обратно после.
    core._config = cfg
    registry.reset_all()
    registry.register(agent_id=agent, core=core)
    try:
        app = create_app(cfg)
        client = TestClient(app)
        resp = client.post(
            "/ingest_document",
            json={
                "agent_id": agent,
                "path": str(FIXTURE_DIR / "sample.txt"),
            },
        )
        assert resp.status_code == 503
        assert "disabled" in resp.json()["detail"]
    finally:
        core.shutdown()
        registry.reset_all()


# ── Whitelist enforcement ───────────────────────────────────────────


def test_ingest_whitelist_blocks_outside(
    clean_db: str, tmp_path: Path
) -> None:
    """Path вне STYX_INGEST_DOC_ROOTS → 422."""
    migrate.run(clean_db)
    inside_root = tmp_path / "allowed"
    inside_root.mkdir()
    outside = tmp_path / "outside"
    outside.mkdir()
    bad = outside / "secret.txt"
    bad.write_text("classified")

    cfg = load_config()
    cfg = replace(
        cfg,
        database_url=clean_db,
        http_token=None,
        ingest_doc_roots=[str(inside_root)],
    )

    agent = "alpha"
    core = StyxMemoryCore(agent_id=agent)
    core._config = cfg
    core.initialize(session_id=str(uuid.uuid4()), agent_identity=agent)
    core._config = cfg  # initialize перезатёр — restore.
    registry.reset_all()
    registry.register(agent_id=agent, core=core)
    try:
        app = create_app(cfg)
        client = TestClient(app)
        resp = client.post(
            "/ingest_document",
            json={"agent_id": agent, "path": str(bad)},
        )
        assert resp.status_code == 422
        assert "allowed roots" in resp.json()["detail"]
    finally:
        core.shutdown()
        registry.reset_all()


def test_ingest_whitelist_allows_inside(
    clean_db: str, tmp_path: Path
) -> None:
    migrate.run(clean_db)
    root = tmp_path / "allowed"
    root.mkdir()
    target = root / "ok.txt"
    target.write_text("Allowed content within whitelist root.")

    cfg = load_config()
    cfg = replace(
        cfg,
        database_url=clean_db,
        http_token=None,
        ingest_doc_roots=[str(root)],
    )

    agent = "alpha"
    core = StyxMemoryCore(agent_id=agent)
    core._config = cfg
    core.initialize(session_id=str(uuid.uuid4()), agent_identity=agent)
    core._config = cfg  # initialize перезатёр — restore.
    registry.reset_all()
    registry.register(agent_id=agent, core=core)
    try:
        app = create_app(cfg)
        client = TestClient(app)
        resp = client.post(
            "/ingest_document",
            json={"agent_id": agent, "path": str(target)},
        )
        assert resp.status_code == 200, resp.text
    finally:
        core.shutdown()
        registry.reset_all()


# ── Size guard ──────────────────────────────────────────────────────


def test_ingest_size_guard(clean_db: str, tmp_path: Path) -> None:
    """File > ingest_doc_max_bytes → 422."""
    migrate.run(clean_db)
    target = tmp_path / "big.txt"
    target.write_bytes(b"x" * 1024)  # 1 KiB

    cfg = load_config()
    cfg = replace(
        cfg,
        database_url=clean_db,
        http_token=None,
        ingest_doc_max_bytes=512,  # 512 байт
    )

    agent = "alpha"
    core = StyxMemoryCore(agent_id=agent)
    core._config = cfg
    core.initialize(session_id=str(uuid.uuid4()), agent_identity=agent)
    core._config = cfg  # initialize перезатёр — restore.
    registry.reset_all()
    registry.register(agent_id=agent, core=core)
    try:
        app = create_app(cfg)
        client = TestClient(app)
        resp = client.post(
            "/ingest_document",
            json={"agent_id": agent, "path": str(target)},
        )
        assert resp.status_code == 422
        assert "too large" in resp.json()["detail"]
    finally:
        core.shutdown()
        registry.reset_all()
