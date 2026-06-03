"""Unit-tests для парсеров документов (волна 28).

Plain text + Markdown тестируются на static fixtures в
`tests/fixtures/documents/`. PDF/DOCX/XLSX — генерируются runtime
через pytest fixtures (session-scoped tmp_path), чтобы не хранить
binary файлы в репо.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from styx.engine.document_parsers import (
    is_supported_extension,
    mime_for_extension,
    normalize_extension,
    parse,
    verify_magic_bytes,
)
from styx.engine.document_parsers._mime import _MAGIC_BYTES


FIXTURE_DIR = Path(__file__).parent.parent.parent / "fixtures" / "documents"


# ── PDF fixture builder ─────────────────────────────────────────────


@pytest.fixture(scope="session")
def pdf_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Сгенерировать минимальный PDF с двумя страницами через pypdf."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    # Создаём страницу простейшего размера; текст добавляется через
    # add_blank_page + add_metadata; для extract_text нужен content
    # stream, проще сгенерировать через reportlab — но reportlab не в
    # deps. Используем pypdf's PdfWriter.add_blank_page + write_stream
    # с simple content stream через `merge_page` это сложно. Простой
    # путь: создаём текстовый PDF через pypdf низкоуровнево.
    writer.add_blank_page(width=200, height=200)
    writer.add_blank_page(width=200, height=200)

    target = tmp_path_factory.mktemp("pdf") / "sample.pdf"
    with target.open("wb") as fh:
        writer.write(fh)
    return target


@pytest.fixture(scope="session")
def pdf_with_text_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """PDF с реальным текстом через pypdf low-level API.

    pypdf не имеет high-level "add text" — генерируем PDF через
    ручной content stream. Простой подход: создаём `PdfWriter`,
    добавляем страницу, и затем write_stream с в content stream
    наносим текст с BT/ET операторами.
    """
    from pypdf import PdfWriter
    from pypdf.generic import (
        ContentStream,
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
        TextStringObject,
    )

    writer = PdfWriter()
    page = writer.add_blank_page(width=595, height=842)

    # Минимальный PDF content stream с текстом "Hello Styx wave 28".
    content_bytes = (
        b"BT /F1 24 Tf 100 700 Td (Hello Styx wave 28) Tj ET"
    )
    stream = DecodedStreamObject()
    stream.set_data(content_bytes)
    page[NameObject("/Contents")] = stream

    # Font resource — стандартный Helvetica (Type1 base14, не требует
    # embedding'а).
    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    resources = page.get("/Resources", DictionaryObject())
    fonts = resources.get("/Font", DictionaryObject())
    fonts[NameObject("/F1")] = font
    resources[NameObject("/Font")] = fonts
    page[NameObject("/Resources")] = resources

    target = tmp_path_factory.mktemp("pdf-text") / "with_text.pdf"
    with target.open("wb") as fh:
        writer.write(fh)
    return target


# ── DOCX fixture builder ────────────────────────────────────────────


@pytest.fixture(scope="session")
def docx_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Сгенерировать минимальный DOCX через python-docx."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Тестовый параграф один.")
    doc.add_paragraph("Второй параграф для проверки.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A1"
    table.rows[0].cells[1].text = "B1"
    table.rows[1].cells[0].text = "A2"
    table.rows[1].cells[1].text = "B2"

    target = tmp_path_factory.mktemp("docx") / "sample.docx"
    doc.save(str(target))
    return target


# ── XLSX fixture builder ────────────────────────────────────────────


@pytest.fixture(scope="session")
def xlsx_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Сгенерировать минимальный XLSX через openpyxl."""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    assert ws1 is not None
    ws1.title = "Sheet1"
    ws1.append(["Имя", "Возраст"])
    ws1.append(["alpha", 28])
    ws1.append(["beta", 30])

    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["X", "Y", "Z"])
    ws2.append([1, 2, 3])

    target = tmp_path_factory.mktemp("xlsx") / "sample.xlsx"
    wb.save(str(target))
    return target


# ── Mime helpers ────────────────────────────────────────────────────


def test_normalize_extension_lowercases() -> None:
    assert normalize_extension(Path("/tmp/Foo.PDF")) == ".pdf"
    assert normalize_extension(Path("/tmp/x.MD")) == ".md"


def test_normalize_extension_no_suffix() -> None:
    assert normalize_extension(Path("/tmp/noext")) == ""


def test_is_supported_extension_known() -> None:
    for ext in (".pdf", ".docx", ".xlsx", ".md", ".markdown", ".txt", ".text"):
        assert is_supported_extension(ext)


def test_is_supported_extension_unknown() -> None:
    assert not is_supported_extension(".pptx")
    assert not is_supported_extension(".rtf")
    assert not is_supported_extension("")


def test_mime_for_extension() -> None:
    assert mime_for_extension(".pdf") == "application/pdf"
    assert (
        mime_for_extension(".docx")
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert mime_for_extension(".md") == "text/markdown"
    assert mime_for_extension(".pptx") is None


def test_verify_magic_bytes_pdf_ok(pdf_fixture: Path) -> None:
    verify_magic_bytes(pdf_fixture, ".pdf")  # no raise


def test_verify_magic_bytes_skips_text() -> None:
    # text/* — magic не проверяем (utf-8 не имеет signature).
    verify_magic_bytes(FIXTURE_DIR / "sample.txt", ".txt")
    verify_magic_bytes(FIXTURE_DIR / "sample.md", ".md")


def test_verify_magic_bytes_mismatch(tmp_path: Path) -> None:
    decoy = tmp_path / "fake.pdf"
    decoy.write_bytes(b"PK\x03\x04not-a-pdf")
    with pytest.raises(ValueError, match="mime mismatch"):
        verify_magic_bytes(decoy, ".pdf")


def test_magic_bytes_table_includes_pdf_docx_xlsx() -> None:
    # Защита от регрессии в таблице (волна 28).
    assert _MAGIC_BYTES[".pdf"] == b"%PDF-"
    assert _MAGIC_BYTES[".docx"] == b"PK\x03\x04"
    assert _MAGIC_BYTES[".xlsx"] == b"PK\x03\x04"


# ── Parser tests ────────────────────────────────────────────────────


def test_parse_pdf_with_text(pdf_with_text_fixture: Path) -> None:
    result = parse(pdf_with_text_fixture)
    assert "Hello Styx wave 28" in result.text
    assert result.metadata.get("page_count") == 1


def test_parse_pdf_blank_pages_empty_text(pdf_fixture: Path) -> None:
    """Blank PDF — extract_text возвращает empty; caller отвечает 422."""
    result = parse(pdf_fixture)
    assert result.text == ""
    assert result.metadata.get("page_count") == 2


def test_parse_docx(docx_fixture: Path) -> None:
    result = parse(docx_fixture)
    assert "Тестовый параграф один." in result.text
    assert "Второй параграф для проверки." in result.text
    assert "A1 | B1" in result.text
    assert "A2 | B2" in result.text
    assert result.metadata.get("paragraph_count", 0) >= 2
    assert result.metadata.get("table_count") == 1


def test_parse_xlsx(xlsx_fixture: Path) -> None:
    result = parse(xlsx_fixture)
    assert "=== Sheet: Sheet1 ===" in result.text
    assert "=== Sheet: Sheet2 ===" in result.text
    assert "alpha" in result.text
    assert "28" in result.text
    assert result.metadata.get("sheet_names") == ["Sheet1", "Sheet2"]


def test_parse_plain_text() -> None:
    result = parse(FIXTURE_DIR / "sample.txt")
    assert "Тестовый plain text документ" in result.text
    assert result.metadata.get("mime_type_input") == "text/plain"
    assert result.metadata.get("line_count", 0) > 1


def test_parse_markdown() -> None:
    result = parse(FIXTURE_DIR / "sample.md")
    assert "# Sample Markdown" in result.text
    assert "```python" in result.text
    assert result.metadata.get("mime_type_input") == "text/markdown"


def test_parse_unsupported_extension(tmp_path: Path) -> None:
    decoy = tmp_path / "foo.pptx"
    decoy.write_bytes(b"PK\x03\x04")
    with pytest.raises(ValueError, match="unsupported extension"):
        parse(decoy)


def test_parse_no_extension(tmp_path: Path) -> None:
    decoy = tmp_path / "noext"
    decoy.write_bytes(b"text")
    with pytest.raises(ValueError, match="unsupported extension"):
        parse(decoy)


def test_parse_pdf_mime_mismatch(tmp_path: Path) -> None:
    decoy = tmp_path / "fake.pdf"
    decoy.write_bytes(b"PK\x03\x04Not really a PDF")
    with pytest.raises(ValueError, match="mime mismatch"):
        parse(decoy)


def test_parse_docx_mime_mismatch(tmp_path: Path) -> None:
    decoy = tmp_path / "fake.docx"
    decoy.write_bytes(b"%PDF-not really docx")
    with pytest.raises(ValueError, match="mime mismatch"):
        parse(decoy)


def test_parse_malformed_pdf(tmp_path: Path) -> None:
    decoy = tmp_path / "malformed.pdf"
    decoy.write_bytes(b"%PDF-1.4\nthis is not a real pdf body")
    with pytest.raises(ValueError, match="PDF"):
        parse(decoy)


def test_parse_malformed_docx(tmp_path: Path) -> None:
    decoy = tmp_path / "malformed.docx"
    decoy.write_bytes(b"PK\x03\x04not really a zip")
    with pytest.raises(ValueError, match="DOCX"):
        parse(decoy)
